import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class FinalReportGenerator:
    def __init__(self, config: dict, eval_result: dict, llm_advisor=None):
        self.doc = Document()
        self.config = config
        self.eval_result = eval_result
        self.llm_advisor = llm_advisor

    def set_font_style(self, run, size_pt=None, bold=False, color=None):
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if size_pt: run.font.size = Pt(size_pt)
        if bold: run.bold = bold
        if color: run.font.color.rgb = color

    def _format_fuzzy(self, fuzzy_dict: dict) -> str:
        if not fuzzy_dict: return "-"
        return ", ".join([f"{k}: {v}%" for k, v in fuzzy_dict.items()])

    def generate(self, fault_cases: list, output_path: str):
        r_info = self.config.get('report_info', {})
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(10.5)

        # 建立中文映射，用于各章节二级标题翻译
        name_mapping = {
            "frame_vibration": "轿架振动",
            "smoothness": "平稳度异常",
            "轿厢健康度": "轿厢平稳度",
            "motor_fault": "电机运行状态",
            "rotor_misalignment": "电机转子不对中",
            "stator_eccentricity": "电机定子偏心",
            "bearing_fault": "轴承组健康度",
            "bolt_loose": "底座刚度防松",
            "water": "水浸异常",
            "temperature": "温度异常",
            "displacement": "物理位移",
            "motor_current": "电流异常",
            "noise_ratio": "异常噪声"
        }

        # 标题
        header = self.doc.add_heading('', 0)
        run = header.add_run(r_info.get('title', '直梯智能化诊断报告'))
        self.set_font_style(run, size_pt=22, bold=True)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 标题1：直梯整体运行健康度 ---
        self.set_font_style(self.doc.add_heading('', level=1).add_run('一、 直梯整体运行健康度'), size_pt=14, bold=True)
        
        sys_score = self.eval_result.get('score', 100)
        sys_fuzzy = self._format_fuzzy(self.eval_result.get('fuzzy_distribution', {}))
        
        p1 = self.doc.add_paragraph()
        self.set_font_style(p1.add_run(f"系统最终健康分: {sys_score} 分 | 整体状态评估: "), bold=True)
        color = RGBColor(255, 0, 0) if sys_score < 75 else RGBColor(0, 150, 0)
        self.set_font_style(p1.add_run(sys_fuzzy), bold=True, color=color)

        table = self.doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
        for cell, text in zip(table.rows[0].cells, ['大类/部件', '综合诊断评估', '检测子项', '子项状态分布']):
            self.set_font_style(cell.paragraphs[0].add_run(text), bold=True)
            
        for device, info in self.eval_result.get('device_scores', {}).items():
            details = info.get('details', {})
            if not details: continue
            
            valid_details = {k: v for k, v in details.items() if isinstance(v, dict) and 'score' in v}
            if not valid_details: continue
            
            start_row_idx = len(table.rows)
            for i, (f_name, detail) in enumerate(valid_details.items()):
                row = table.add_row().cells
                if i == 0:
                    row[0].text = device
                    row[1].text = f"{info['score']} 分\n[{self._format_fuzzy(info.get('fuzzy_distribution', {}))}]"
                
                display_name = name_mapping.get(f_name, f_name)
                cell_2_text = display_name
                cell_3_text = f"{detail['score']} 分\n[{self._format_fuzzy(detail.get('fuzzy_distribution', {}))}]"
                
                if 'sub_details' in detail:
                    for sub_name, sub_info in detail['sub_details'].items():
                        cell_2_text += f"\n  └─ {sub_name}"
                        cell_3_text += f"\n  └─ {sub_info['score']} 分 [{self._format_fuzzy(sub_info.get('fuzzy_distribution', {}))}]"
                
                row[2].text = cell_2_text
                row[3].text = cell_3_text
            
            if len(valid_details) > 1:
                end_row_idx = len(table.rows) - 1
                table.cell(start_row_idx, 0).merge(table.cell(end_row_idx, 0))
                table.cell(start_row_idx, 1).merge(table.cell(end_row_idx, 1))

        if self.eval_result.get('env_scores'):
            env_start = len(table.rows)
            for i, (env_name, info) in enumerate(self.eval_result['env_scores'].items()):
                row = table.add_row().cells
                if i == 0: row[0].text, row[1].text = "环境与电气检测", "熔断否决机制"
                row[2].text = name_mapping.get(env_name, env_name)
                row[3].text = f"{info['score']} 分\n[{self._format_fuzzy(info.get('fuzzy_distribution', {}))}]"
            if len(self.eval_result['env_scores']) > 1:
                table.cell(env_start, 0).merge(table.cell(len(table.rows) - 1, 0))
                table.cell(env_start, 1).merge(table.cell(len(table.rows) - 1, 1))

        # --- 标题2：设备环境与采样信息 (按子项切分二级标题) ---
        self.set_font_style(self.doc.add_heading('', level=1).add_run('二、 报警源定位信息'), size_pt=14, bold=True)
        if not fault_cases:
            self.doc.add_paragraph().add_run("系统无报警源，未定位到异常设备。")
        else:
            for idx, case in enumerate(fault_cases, 1):
                device_name = case["device"]
                fault_name = case["fault_name"]
                disp_name = name_mapping.get(fault_name, fault_name)
                data_info = case["data_info"]
                
                # 二级标题
                self.set_font_style(self.doc.add_heading('', level=2).add_run(f'2.{idx} 【{device_name}】 {disp_name}'), size_pt=12, bold=True)
                
                self.doc.add_paragraph().add_run(r_info.get('intro_text', '').format(
                    loc=data_info.get('location', '未知'), comp=data_info.get('component', '未知')))
                
                info_table = self.doc.add_table(rows=3, cols=2); info_table.style = 'Table Grid'
                for i, (k, v) in enumerate([("传感器编号", data_info.get('sensor_id', 'N/A')), 
                                            ("采样率(Hz)", str(data_info.get('fs', 'N/A'))), 
                                            ("通道数", str(len(data_info.get('signal_data', []))))]):
                    info_table.rows[i].cells[0].text, info_table.rows[i].cells[1].text = k, v

        # --- 标题3：深度诊断结论 ---
        self.set_font_style(self.doc.add_heading('', level=1).add_run('三、 深度诊断结论'), size_pt=14, bold=True)
        if not fault_cases:
            run_res = self.doc.add_paragraph().add_run("判定结果：系统整体运行优良，未发现明显风险隐患。")
            self.set_font_style(run_res, size_pt=14, bold=True, color=RGBColor(0, 150, 0))
        else:
            for idx, case in enumerate(fault_cases, 1):
                device_name = case["device"]
                fault_name = case["fault_name"]
                disp_name = name_mapping.get(fault_name, fault_name)
                detail = case["detail"]
                
                self.set_font_style(self.doc.add_heading('', level=2).add_run(f'3.{idx} 结论项目：{disp_name}'), size_pt=12, bold=True)
                
                p3 = self.doc.add_paragraph()
                run_res = p3.add_run(r_info.get('conclusion_text', '').format(device=device_name, fault=disp_name))
                self.set_font_style(run_res, size_pt=13, bold=True, color=RGBColor(255, 0, 0))
                
                fuzzy_str = self._format_fuzzy(detail.get('fuzzy_distribution', {}))
                self.doc.add_paragraph().add_run(r_info.get('desc_text', '').format(score=detail.get('score', 'N/A'), fuzzy_str=fuzzy_str))

        # --- 标题4：可视化分析验证 ---
        self.set_font_style(self.doc.add_heading('', level=1).add_run('四、 可视化分析验证'), size_pt=14, bold=True)
        if not fault_cases:
            self.set_font_style(self.doc.add_paragraph().add_run("系统运行优良，无异常波形需验证。"), color=RGBColor(128, 128, 128))
        else:
            for idx, case in enumerate(fault_cases, 1):
                device_name = case["device"]
                fault_name = case["fault_name"]
                disp_name = name_mapping.get(fault_name, fault_name)
                img_paths = case["img_paths"]
                
                self.set_font_style(self.doc.add_heading('', level=2).add_run(f'4.{idx} 【{device_name}】 {disp_name} 波形特征'), size_pt=12, bold=True)
                
                if img_paths:
                    img_titles = [('time', f'图{idx}-1：时域波形'), ('spectrum', f'图{idx}-2：频谱图 (标注特征谐波)'), ('envelope', f'图{idx}-3：包络解调谱 (识别频段缺陷)')]
                    for key, title in img_titles:
                        if key in img_paths and os.path.exists(img_paths[key]):
                            self.doc.add_picture(img_paths[key], width=Inches(6.0))
                            self.doc.add_paragraph(title).alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    fallback_text = r_info.get('no_img_text', '') if device_name == "环境与电气" else "系统未能从缓存中提取到该子项的有效波形数据。"
                    self.set_font_style(self.doc.add_paragraph().add_run(fallback_text), color=RGBColor(128, 128, 128))

        # --- 标题5：专家评估与维修建议 ---
        self.set_font_style(self.doc.add_heading('', level=1).add_run('五、 专家多级定性建议'), size_pt=14, bold=True)
        if not fault_cases:
            self.doc.add_paragraph().add_run("【良好】各项运行指标处于健康状态，建议保持常规巡检计划，无需特殊干预。")
        else:
            for idx, case in enumerate(fault_cases, 1):
                fault_name = case["fault_name"]
                detail = case["detail"]
                disp_name = name_mapping.get(fault_name, fault_name)
                
                self.set_font_style(self.doc.add_heading('', level=2).add_run(f'5.{idx} 针对 [{disp_name}] 的行动建议'), size_pt=12, bold=True)
                
                fuzzy_dist = detail.get('fuzzy_distribution', {})
                primary_grade = max(fuzzy_dist, key=fuzzy_dist.get) if fuzzy_dist else 'H4'

                advice_key = 'bearing_fault' if fault_name.startswith('bearing_') and fault_name != 'bearing_cage' else fault_name

                if self.llm_advisor is not None:
                    advice_text = self.llm_advisor.generate_advice(
                        fault_type=advice_key,
                        risk_grade=primary_grade,
                        score=detail.get('score'),
                        fuzzy_dist=fuzzy_dist,
                    )
                else:
                    target_advice = self.config.get('expert_advice', {}).get(advice_key, self.config.get('expert_advice', {}).get('default', {}))
                    advice_text = target_advice.get(primary_grade, "暂无特定级别建议，请结合现场物理排查确认。")

                self.doc.add_paragraph().add_run(advice_text)

        # --- 标题6：标准免责及报告说明 ---
        self.set_font_style(self.doc.add_heading('', level=1).add_run('六、 报告说明'), size_pt=14, bold=True)
        disclaimer = self.config.get('report_info', {}).get(
            'disclaimer_text', 
            '本报告由系统自动生成，结果仅供参考，不作为法定依据。'
        )
        p_disclaimer = self.doc.add_paragraph()
        run_disclaimer = p_disclaimer.add_run(disclaimer)
        self.set_font_style(run_disclaimer, size_pt=10.5, color=RGBColor(128, 128, 128))
        p_disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.save(output_path)
        print(f"\n[+] 诊断结果已成功归档至: {output_path}")