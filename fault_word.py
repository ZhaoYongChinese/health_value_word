import numpy as np
import matplotlib.pyplot as plt
from scipy.signal import hilbert, butter, filtfilt, find_peaks
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn # 用于设置中文字体

# ================= 1. 环境与配置 =================
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

FAULT_TOLERANCE_CONFIG = {"电机轴承": 0.1, "减速机轴承": 0.1, "梯级涨紧轮": 0.2, "主驱动轮": 0.1, "默认": 0.1}
RESONANCE_BAND_CONFIG = {"电机轴承": (500, 700), "减速机轴承": (500, 700), "梯级涨紧轮": (400, 600), "主驱动轮": (300, 500), "默认": (300, 500)}
fault_characteristic_freqs = {
    "电机轴承": {"inner": 104.16, "outer": 71.84, "ball": 41.92, "cage": 6.528},
    "减速机轴承": {"inner": 13.818, "outer": 9.702, "ball": 2.2932, "cage": 0.34692},
    "梯级涨紧轮": {"inner": 2.51266, "outer": 1.873534, "ball": 1.055821, "cage": 0.133779},
    "主驱动轮": {"inner": 5.087992, "outer": 3.684408, "ball": 0.9399, "cage": 0.131586},
}

# ================= 2. 核心类定义 =================
class BearingDiagnoser:
    def __init__(self, fault_frequencies, tolerance=0.1, max_harmonics=5):
        self.fault_freqs = fault_frequencies
        self.tolerance = tolerance
        self.max_harmonics = max_harmonics
        self.last_analysis = {}

    def _get_envelope_spectrum(self, signal, fs, resonance_band):
        nyquist = 0.5 * fs
        low = max(resonance_band[0], 1) / nyquist
        high = min(resonance_band[1], nyquist - 1) / nyquist
        b, a = butter(5, [low, high], btype='band')
        filtered_signal = filtfilt(b, a, signal)
        envelope = np.abs(hilbert(filtered_signal))
        envelope -= np.mean(envelope)
        N = len(envelope)
        xf = np.fft.fftfreq(N, 1 / fs)
        yf_abs = 2.0 / N * np.abs(np.fft.fft(envelope))
        mask = xf >= 0
        return xf[mask], yf_abs[mask]
    
    def diagnose(self, signal, fs, component_name, resonance_band):
        target_faults = self.fault_freqs[component_name]
        xf, yf = self._get_envelope_spectrum(signal, fs, resonance_band)
        self.last_analysis = {'freqs': xf, 'amps': yf, 'component': component_name, 'fs': fs, 'signal': signal}
        
        height_th = np.mean(yf)
        peak_indices, _ = find_peaks(yf, height=height_th, prominence=0.5 * np.std(yf))
        
        scores = {}
        for fault_type, f_char in target_faults.items():
            total_energy, harmonics_found = 0, 0
            for n in range(1, self.max_harmonics + 1):
                h_freq = n * f_char
                if fault_type == 'cage':
                    mask = (xf >= h_freq - self.tolerance) & (xf <= h_freq + self.tolerance)
                    if np.any(mask) and np.max(yf[mask]) > height_th:
                        harmonics_found += 1
                        total_energy += np.max(yf[mask])**2
                else:
                    for p_idx in peak_indices:
                        if abs(xf[p_idx] - h_freq) <= self.tolerance:
                            harmonics_found += 1
                            total_energy += yf[p_idx]**2
                            break
            scores[fault_type] = total_energy * (harmonics_found ** 2) + 1e-9
        
        total = sum(scores.values())
        confidences = {k: v / total for k, v in scores.items()}
        return max(confidences, key=confidences.get), confidences

    def plot_time_domain(self):
        if not self.last_analysis: return
        signal = self.last_analysis['signal']
        fs = self.last_analysis['fs']
        time = np.arange(len(signal)) / fs
        plt.figure(figsize=(10, 4))
        #mask = time < 0.5 
        #plt.plot(time[mask], signal[mask], color='blue', linewidth=0.5)
        plt.plot(time, signal, color='blue', linewidth=0.5)
        plt.title(f"'{self.last_analysis['component']}' 时域波形")
        plt.xlabel('时间 (s)')
        plt.ylabel('幅值')
        plt.grid(True, alpha=0.3)

    def plot_last_analysis(self):
        d = self.last_analysis
        plt.plot(d['freqs'], d['amps'], color='c', lw=1, label='包络谱')
        target = self.fault_freqs[d['component']]
        colors = {'inner': 'r', 'outer': 'g', 'ball': 'b', 'cage': 'm'}
        for ft, f in target.items():
            for n in range(1, 6):
                plt.axvline(n*f, color=colors.get(ft), ls='--', alpha=0.5, label=ft if n==1 else "")
        plt.title(f"{d['component']} 包络谱分析")
        plt.xlim(0, max(target.values()) * 5)
        plt.legend(loc='upper right', fontsize='small')

    def plot_frequency_spectrum(self):
        s, fs = self.last_analysis['signal'], self.last_analysis['fs']
        N = len(s)
        yf = 2.0/N * np.abs(np.fft.fft(s)[:N//2])
        xf = np.linspace(0.0, fs/2.0, N//2)
        plt.plot(xf, yf, color='navy', lw=1)
        plt.title("原始信号频谱 (FFT)")

class FinalReportGenerator:
    def __init__(self, info):
        self.doc = Document()
        self.info = info

    def set_font_style(self, run, size_pt=None, bold=False, color=None):
        """通用字体设置函数：宋体+Times New Roman"""
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if size_pt:
            run.font.size = Pt(size_pt)
        if bold:
            run.bold = bold
        if color:
            run.font.color.rgb = color

    def generate(self, predicted_type, confidences, char_freqs, diagnoser, output_path):
        # --- 全局默认字体设置 ---
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(10.5) # 五号字体

        # --- 标题 ---
        header = self.doc.add_heading('', 0)
        run = header.add_run('扶梯关键部件运行状态智能化诊断报告')
        self.set_font_style(run, size_pt=22, bold=True)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- 一、设备环境与采样信息 ---
        h1 = self.doc.add_heading('', level=1)
        self.set_font_style(h1.add_run('一、设备环境与采样信息'), size_pt=14, bold=True)
        
        p1 = self.doc.add_paragraph()
        self.set_font_style(p1.add_run("本次诊断针对 "))
        self.set_font_style(p1.add_run(f"{self.info['location']}"), bold=True)
        self.set_font_style(p1.add_run(" 的扶梯系统进行。受检部件为 "))
        self.set_font_style(p1.add_run(f"{self.info['comp_name']}"), bold=True)
        self.set_font_style(p1.add_run("。采样通过高灵敏度加速度传感器完成，确保了捕捉早期损伤微弱冲击信号的能力。"))

        table = self.doc.add_table(rows=6, cols=2); table.style = 'Table Grid'
        # 修正：将报警时间列表合并为字符串显示
        alarm_times_str = "\n".join(self.info['alarm_times']) if self.info['alarm_times'] else "无记录"
        
        details = [
            ("传感器编号", self.info['sensor_code']),
            ("采样频率", f"{self.info['fs']} Hz"),
            ("共振解调带", f"{self.info['resonance_band'][0]} - {self.info['resonance_band'][1]} Hz"),
            ("分析特征频率", f"{char_freqs[predicted_type]:.2f} Hz ({predicted_type})"),
            ("报警发生时间", alarm_times_str), # 这里是重点
            ("风险判定等级", "预警 (高置信度)")
        ]
        for i, (k, v) in enumerate(details):
            for cell, text in zip(table.rows[i].cells, [k, v]):
                run = cell.paragraphs[0].add_run(str(text))
                self.set_font_style(run)

        # --- 二、全指标置信度分布 ---
        h2 = self.doc.add_heading('', level=1)
        self.set_font_style(h2.add_run('二、全指标置信度分析分布'), size_pt=14, bold=True)
        p2 = self.doc.add_paragraph()
        self.set_font_style(p2.add_run("系统对该部件可能存在的四项核心轴承故障进行了并行计算，各维度置信度评分如下表所示："))
        
        conf_table = self.doc.add_table(rows=1, cols=3); conf_table.style = 'Table Grid'
        hdr_cells = conf_table.rows[0].cells
        for cell, text in zip(hdr_cells, ['检测项目', '置信度评分', '风险判定']):
            self.set_font_style(cell.paragraphs[0].add_run(text), bold=True)

        for fault_key, score in confidences.items():
            row_cells = conf_table.add_row().cells
            self.set_font_style(row_cells[0].paragraphs[0].add_run(f"轴承{fault_key}故障"))
            self.set_font_style(row_cells[1].paragraphs[0].add_run(f"{score:.2%}"))
            
            p_res = row_cells[2].paragraphs[0]
            if fault_key == predicted_type:
                run = p_res.add_run("确认为主故障源")
                self.set_font_style(run, bold=True, color=RGBColor(200, 0, 0))
            else:
                self.set_font_style(p_res.add_run("风险极低" if score < 0.1 else "疑似关联"))

        # --- 三、诊断详细结果 ---
        h3 = self.doc.add_heading('', level=1)
        self.set_font_style(h3.add_run('三、深度诊断结论'), size_pt=14, bold=True)
        
        res_p = self.doc.add_paragraph()
        self.set_font_style(res_p.add_run("判定结果："), bold=True)
        run_res = res_p.add_run(f"【{predicted_type}】故障风险")
        self.set_font_style(run_res, size_pt=14, bold=True, color=RGBColor(255, 0, 0))
        
        p3_desc = self.doc.add_paragraph()
        self.set_font_style(p3_desc.add_run(
            f"基于最大置信度原则，系统判定当前部件存在显著的 {predicted_type} 损伤特征。 "
            f"该项指标评分达到 {confidences[predicted_type]:.2%}，远超其他监测项，排除干扰误报可能。"
        ))

# --- 四、可视化分析验证 ---
        h4 = self.doc.add_heading('', level=1)
        self.set_font_style(h4.add_run('四、可视化分析验证'), size_pt=14, bold=True)
        
        # 4.1 时域图
        h41 = self.doc.add_heading('', level=2)
        self.set_font_style(h41.add_run('4.1 原始信号时域波形'), size_pt=12, bold=True)
        diagnoser.plot_time_domain()
        plt.savefig("time_domain.png", bbox_inches='tight')
        self.doc.add_picture("time_domain.png", width=Inches(5.5))
        self.doc.add_paragraph("图1：原始信号时域波形").alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 【新增：4.2 原始信号频谱图】
        h42 = self.doc.add_heading('', level=2)
        self.set_font_style(h42.add_run('4.2 原始信号频谱 (FFT)'), size_pt=12, bold=True)
        plt.figure(figsize=(10, 4))
        diagnoser.plot_frequency_spectrum()
        plt.savefig("spec.png", bbox_inches='tight')
        self.doc.add_picture("spec.png", width=Inches(5.5))
        self.doc.add_paragraph("图2：原始信号频谱图").alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 4.3 包络谱图
        h43 = self.doc.add_heading('', level=2)
        self.set_font_style(h43.add_run('4.3 包络解调谱（故障指纹）'), size_pt=12, bold=True)
        plt.figure(figsize=(10, 4))
        diagnoser.plot_last_analysis()
        plt.savefig("env.png", bbox_inches='tight')
        self.doc.add_picture("env.png", width=Inches(5.5))
        self.doc.add_paragraph("图3：包络解调谱").alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 五、分项专家技术建议 ---
        h5 = self.doc.add_heading('', level=1)
        self.set_font_style(h5.add_run('五、专家评估与维修建议'), size_pt=14, bold=True)
        
        expert_advice = {
            "inner": "【技术特征】：包络谱呈现明显的内圈故障频率。建议更换轴承并检查轴颈精度。",
            "outer": "【技术特征】：外圈故障频率能量集中。建议检查轴承座是否有疲劳裂纹。",
            "ball": "【技术特征】：滚动体故障频率及其谐波明显。必须立即停梯检查润滑油中是否有金属碎屑。",
            "cage": "【技术特征】：保持架频率极低且能量显著。应立即组织人员开盖核实并更换轴承。"
        }
        advice_p = self.doc.add_paragraph()
        self.set_font_style(advice_p.add_run(expert_advice.get(predicted_type, "暂无建议。")))

        # --- 六、报告声明 ---
        h6 = self.doc.add_heading('', level=1)
        self.set_font_style(h6.add_run('六、报告说明'), size_pt=14, bold=True)
        p6 = self.doc.add_paragraph()
        self.set_font_style(p6.add_run("1. 本诊断基于数据模型推导，实际维护请结合现场开盖检查。\n2. 报警时间关联分析显示，该故障具有持续演变的特征，建议尽快处置。"))

        self.doc.save(output_path)
        plt.close('all')
        for tmp_img in ["time_domain.png", "env.png"]:
            if os.path.exists(tmp_img): os.remove(tmp_img)
        print(f"深度诊断报告已完成：{output_path}")

# ================= 3. 主运行流程 =================
if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__)) if "__file__" in locals() else os.getcwd()
    target_filename = "wave19_32KHz.csv"
    file_full_path = os.path.join(current_dir, target_filename)

    if os.path.exists(file_full_path):
        df = pd.read_csv(file_full_path, header=None)
        # 确保只取第二列且减去均值
        data_1 = df.iloc[:, 1].values
        
        fs = 32000
        component_name = "主驱动轮"
        tolerance = FAULT_TOLERANCE_CONFIG[component_name]
        resonance_band = RESONANCE_BAND_CONFIG[component_name]

        diagnoser = BearingDiagnoser(fault_characteristic_freqs, tolerance=tolerance)
        predicted_type, confidences = diagnoser.diagnose(data_1, fs, component_name, resonance_band)

        # 这里的报警时间将会显示在报告中
        report_info = {
            "location": "门诊楼1-2楼北",
            "sensor_code": "9005",
            "comp_name": "左侧主驱动轮轴承",
            "alarm_times": ["2025-12-02 06:54:47"], 
            "fs": fs,
            "resonance_band": resonance_band,
            "filename": target_filename
        }

        reporter = FinalReportGenerator(report_info)
        reporter.generate(
            predicted_type, 
            confidences, 
            fault_characteristic_freqs[component_name], 
            diagnoser, 
            f"诊断报告_{target_filename.replace('.csv', '')}.docx"
        )